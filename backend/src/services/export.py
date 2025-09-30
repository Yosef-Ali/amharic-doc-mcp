"""Export service for generating various document formats."""

from __future__ import annotations

import json
import logging
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from uuid import UUID

from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from ..db.models.document import Document
from ..db.models.export_template import ExportTemplate, ExportFormat
from ..config.settings import Settings

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting documents in various formats."""
    
    def __init__(self, settings: Settings):
        self.settings = settings
        
    async def export_document(
        self,
        session: AsyncSession,
        document_id: UUID,
        export_format: ExportFormat,
        template_id: Optional[UUID] = None,
        options: Optional[Dict[str, Any]] = None
    ) -> Tuple[bytes, str]:
        """Export a document in the specified format."""
        # Get document
        result = await session.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        
        if not document:
            raise ValueError(f"Document {document_id} not found")
            
        # Get extracted content (would normally come from MongoDB)
        extracted_content = await self._get_extracted_content(document_id)
        
        # Get export template if specified
        template = None
        if template_id:
            result = await session.execute(
                select(ExportTemplate).where(ExportTemplate.id == template_id)
            )
            template = result.scalar_one_or_none()
            
        # Generate export based on format
        if export_format == ExportFormat.PDF:
            return await self._export_pdf(document, extracted_content, template, options)
        elif export_format == ExportFormat.DOCX:
            return await self._export_docx(document, extracted_content, template, options)
        elif export_format == ExportFormat.HTML:
            return await self._export_html(document, extracted_content, template, options)
        elif export_format == ExportFormat.MARKDOWN:
            return await self._export_markdown(document, extracted_content, template, options)
        elif export_format == ExportFormat.JSON:
            return await self._export_json(document, extracted_content, template, options)
        else:
            raise ValueError(f"Unsupported export format: {export_format}")
            
    async def create_export_template(
        self,
        session: AsyncSession,
        name: str,
        format: ExportFormat,
        template_content: Dict[str, Any],
        description: Optional[str] = None
    ) -> ExportTemplate:
        """Create a new export template."""
        template = ExportTemplate(
            name=name,
            format=format,
            template_content=template_content,
            description=description,
            created_at=datetime.utcnow()
        )
        
        session.add(template)
        await session.commit()
        await session.refresh(template)
        
        logger.info(f"Created export template '{name}' for format {format}")
        return template
        
    async def get_export_templates(
        self,
        session: AsyncSession,
        format: Optional[ExportFormat] = None
    ) -> List[ExportTemplate]:
        """Get export templates, optionally filtered by format."""
        query = select(ExportTemplate)
        
        if format:
            query = query.where(ExportTemplate.format == format)
            
        result = await session.execute(query)
        return result.scalars().all()
        
    async def delete_export_template(
        self,
        session: AsyncSession,
        template_id: UUID
    ) -> bool:
        """Delete an export template."""
        result = await session.execute(
            select(ExportTemplate).where(ExportTemplate.id == template_id)
        )
        template = result.scalar_one_or_none()
        
        if template:
            await session.delete(template)
            await session.commit()
            logger.info(f"Deleted export template {template_id}")
            return True
            
        return False
        
    async def _get_extracted_content(self, document_id: UUID) -> Dict[str, Any]:
        """Get extracted content for a document (from MongoDB)."""
        # This would normally fetch from MongoDB
        # For now, returning mock data
        return {
            "text": "የአማርኛ ሰነድ ናሙና ይዘት። This is sample Amharic document content.",
            "language": "amh",
            "pages": [
                {
                    "page_number": 1,
                    "text": "የአማርኛ ሰነድ ናሙና ይዘት።",
                    "confidence": 0.95
                }
            ],
            "metadata": {
                "author": "Sample Author",
                "creation_date": "2024-01-01",
                "subject": "Sample Document"
            },
            "entities": [
                {"text": "አማርኛ", "label": "LANGUAGE", "confidence": 0.98}
            ],
            "summary": "የአማርኛ ሰነድ ማጠቃለያ"
        }
        
    async def _export_pdf(
        self,
        document: Document,
        content: Dict[str, Any],
        template: Optional[ExportTemplate],
        options: Optional[Dict[str, Any]]
    ) -> Tuple[bytes, str]:
        """Export document as PDF."""
        buffer = BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom style for Amharic text
        amharic_style = ParagraphStyle(
            'AmharicStyle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',  # Would use Amharic font in production
            fontSize=12,
            spaceAfter=12,
            alignment=0  # Left alignment
        )
        
        # Build story
        story = []
        
        # Title
        title = Paragraph(f"Document: {document.filename}", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Metadata
        if "metadata" in content:
            metadata = content["metadata"]
            for key, value in metadata.items():
                meta_para = Paragraph(f"<b>{key.title()}:</b> {value}", styles['Normal'])
                story.append(meta_para)
                
        story.append(Spacer(1, 12))
        
        # Content
        if "text" in content:
            text_para = Paragraph(content["text"], amharic_style)
            story.append(text_para)
            
        # Summary if available
        if "summary" in content:
            story.append(Spacer(1, 12))
            summary_title = Paragraph("<b>Summary / ማጠቃለያ:</b>", styles['Heading2'])
            story.append(summary_title)
            summary_para = Paragraph(content["summary"], amharic_style)
            story.append(summary_para)
            
        # Add watermark if specified
        if options and options.get("add_watermark"):
            # This would add a watermark to each page
            pass
            
        # Build PDF
        doc.build(story)
        
        # Get PDF bytes
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        filename = f"{document.filename}_export.pdf"
        return pdf_bytes, filename
        
    async def _export_docx(
        self,
        document: Document,
        content: Dict[str, Any],
        template: Optional[ExportTemplate],
        options: Optional[Dict[str, Any]]
    ) -> Tuple[bytes, str]:
        """Export document as DOCX."""
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches
        except ImportError:
            raise ImportError("python-docx is required for DOCX export")
            
        # Create new document
        doc = DocxDocument()
        
        # Add title
        title = doc.add_heading(f"Document: {document.filename}", 0)
        
        # Add metadata
        if "metadata" in content:
            doc.add_heading('Metadata', level=1)
            metadata = content["metadata"]
            for key, value in metadata.items():
                paragraph = doc.add_paragraph()
                paragraph.add_run(f"{key.title()}: ").bold = True
                paragraph.add_run(str(value))
                
        # Add content
        if "text" in content:
            doc.add_heading('Content', level=1)
            doc.add_paragraph(content["text"])
            
        # Add summary
        if "summary" in content:
            doc.add_heading('Summary / ማጠቃለያ', level=1)
            doc.add_paragraph(content["summary"])
            
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        
        docx_bytes = buffer.getvalue()
        buffer.close()
        
        filename = f"{document.filename}_export.docx"
        return docx_bytes, filename
        
    async def _export_html(
        self,
        document: Document,
        content: Dict[str, Any],
        template: Optional[ExportTemplate],
        options: Optional[Dict[str, Any]]
    ) -> Tuple[bytes, str]:
        """Export document as HTML."""
        html_content = f"""
<!DOCTYPE html>
<html lang="am">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document: {document.filename}</title>
    <style>
        body {{
            font-family: 'Noto Sans Ethiopic', sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        .metadata {{
            background-color: #f5f5f5;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
        }}
        .content {{
            text-align: justify;
            margin-bottom: 20px;
        }}
        .summary {{
            border-left: 4px solid #007bff;
            padding-left: 15px;
            background-color: #f8f9fa;
            padding: 15px;
        }}
        h1, h2 {{
            color: #333;
        }}
    </style>
</head>
<body>
    <h1>Document: {document.filename}</h1>
"""
        
        # Add metadata
        if "metadata" in content:
            html_content += "<div class='metadata'><h2>Metadata</h2>"
            metadata = content["metadata"]
            for key, value in metadata.items():
                html_content += f"<p><strong>{key.title()}:</strong> {value}</p>"
            html_content += "</div>"
            
        # Add content
        if "text" in content:
            html_content += f"<div class='content'><h2>Content</h2><p>{content['text']}</p></div>"
            
        # Add summary
        if "summary" in content:
            html_content += f"<div class='summary'><h2>Summary / ማጠቃለያ</h2><p>{content['summary']}</p></div>"
            
        html_content += """
</body>
</html>"""
        
        html_bytes = html_content.encode('utf-8')
        filename = f"{document.filename}_export.html"
        return html_bytes, filename
        
    async def _export_markdown(
        self,
        document: Document,
        content: Dict[str, Any],
        template: Optional[ExportTemplate],
        options: Optional[Dict[str, Any]]
    ) -> Tuple[bytes, str]:
        """Export document as Markdown."""
        markdown_content = f"# Document: {document.filename}\n\n"
        
        # Add metadata
        if "metadata" in content:
            markdown_content += "## Metadata\n\n"
            metadata = content["metadata"]
            for key, value in metadata.items():
                markdown_content += f"**{key.title()}:** {value}\n\n"
                
        # Add content
        if "text" in content:
            markdown_content += "## Content\n\n"
            markdown_content += f"{content['text']}\n\n"
            
        # Add summary
        if "summary" in content:
            markdown_content += "## Summary / ማጠቃለያ\n\n"
            markdown_content += f"{content['summary']}\n\n"
            
        markdown_bytes = markdown_content.encode('utf-8')
        filename = f"{document.filename}_export.md"
        return markdown_bytes, filename
        
    async def _export_json(
        self,
        document: Document,
        content: Dict[str, Any],
        template: Optional[ExportTemplate],
        options: Optional[Dict[str, Any]]
    ) -> Tuple[bytes, str]:
        """Export document as JSON."""
        export_data = {
            "document": {
                "id": str(document.id),
                "filename": document.filename,
                "content_type": document.content_type,
                "file_size": document.file_size,
                "created_at": document.created_at.isoformat(),
                "updated_at": document.updated_at.isoformat() if document.updated_at else None,
                "status": str(document.status),
                "metadata": document.metadata
            },
            "extracted_content": content,
            "export_info": {
                "exported_at": datetime.utcnow().isoformat(),
                "export_format": "JSON",
                "template_used": str(template.id) if template else None
            }
        }
        
        json_content = json.dumps(export_data, ensure_ascii=False, indent=2)
        json_bytes = json_content.encode('utf-8')
        filename = f"{document.filename}_export.json"
        return json_bytes, filename


# Global export service instance
_export_service: Optional[ExportService] = None


def get_export_service(settings: Settings) -> ExportService:
    """Get the global export service instance."""
    global _export_service
    if _export_service is None:
        _export_service = ExportService(settings)
    return _export_service