"""Word document extractor agent - simplified implementation."""

from __future__ import annotations
import logging
from typing import Dict, Any
from ...config.settings import Settings

logger = logging.getLogger(__name__)

class WordExtractorAgent:
    def __init__(self, settings: Settings):
        self.settings = settings
    
    async def extract_word_content(self, file_path: str, document_id) -> Dict[str, Any]:
        """Extract content from Word documents."""
        try:
            from docx import Document
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "success": True,
                "extracted_text": text,
                "paragraph_count": len(doc.paragraphs),
                "processing_method": "docx_extraction"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

def get_word_extractor(settings: Settings) -> WordExtractorAgent:
    return WordExtractorAgent(settings)