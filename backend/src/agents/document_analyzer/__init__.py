"""Document analyzer agent for detecting document types and analyzing structure."""

from __future__ import annotations

import logging
import mimetypes
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from uuid import UUID

from PIL import Image
import cv2
import numpy as np
from sqlalchemy.ext.asyncio import AsyncSession

from ...db.models.document import Document
from ...config.settings import Settings

logger = logging.getLogger(__name__)


class DocumentAnalyzerAgent:
    """Agent for analyzing document structure, type, and processing requirements."""
    
    def __init__(self, settings: Settings):
        self.settings = settings
        self.supported_formats = {
            'image': ['.jpg', '.jpeg', '.png', '.tiff', '.bmp'],
            'pdf': ['.pdf'],
            'word': ['.docx', '.doc'],
            'csv': ['.csv'],
            'text': ['.txt']
        }
        
    async def analyze_document(
        self,
        session: AsyncSession,
        document: Document,
        file_path: str
    ) -> Dict[str, Any]:
        """Comprehensive document analysis."""
        logger.info(f"Starting analysis for document {document.id}")
        
        try:
            analysis_result = {
                "document_id": str(document.id),
                "filename": document.filename,
                "file_size": document.file_size,
                "content_type": document.content_type,
                "analysis_timestamp": document.created_at.isoformat(),
                "basic_info": {},
                "structure_info": {},
                "processing_recommendations": {},
                "confidence_scores": {},
                "detected_issues": []
            }
            
            # Basic file analysis
            basic_info = await self._analyze_basic_properties(file_path, document)
            analysis_result["basic_info"] = basic_info
            
            # Determine document type and format
            doc_type = self._determine_document_type(document.filename, document.content_type)
            analysis_result["document_type"] = doc_type
            
            # Structure analysis based on type
            if doc_type == "image":
                structure_info = await self._analyze_image_structure(file_path)
            elif doc_type == "pdf":
                structure_info = await self._analyze_pdf_structure(file_path)
            elif doc_type == "word":
                structure_info = await self._analyze_word_structure(file_path)
            elif doc_type == "csv":
                structure_info = await self._analyze_csv_structure(file_path)
            else:
                structure_info = {"type": "unknown", "analysis_limited": True}
                
            analysis_result["structure_info"] = structure_info
            
            # Generate processing recommendations
            recommendations = self._generate_processing_recommendations(
                doc_type, structure_info, basic_info
            )
            analysis_result["processing_recommendations"] = recommendations
            
            # Calculate confidence scores
            confidence_scores = self._calculate_confidence_scores(
                doc_type, structure_info, basic_info
            )
            analysis_result["confidence_scores"] = confidence_scores
            
            # Detect potential issues
            issues = self._detect_processing_issues(
                doc_type, structure_info, basic_info
            )
            analysis_result["detected_issues"] = issues
            
            logger.info(f"Document analysis completed for {document.id}")
            
            return {
                "success": True,
                "analysis": analysis_result,
                "processing_complexity": recommendations.get("complexity", "medium"),
                "estimated_processing_time": recommendations.get("estimated_time", 30)
            }
            
        except Exception as e:
            logger.error(f"Document analysis failed for {document.id}: {e}")
            return {
                "success": False,
                "error": str(e),
                "document_id": str(document.id)
            }
            
    async def _analyze_basic_properties(
        self,
        file_path: str,
        document: Document
    ) -> Dict[str, Any]:
        """Analyze basic file properties."""
        path = Path(file_path)
        
        basic_info = {
            "file_exists": path.exists(),
            "file_extension": path.suffix.lower(),
            "mime_type": mimetypes.guess_type(file_path)[0],
            "file_size_mb": round(document.file_size / (1024 * 1024), 2) if document.file_size else 0,
            "is_readable": path.is_file() and path.stat().st_size > 0 if path.exists() else False
        }
        
        return basic_info
        
    def _determine_document_type(self, filename: str, content_type: str) -> str:
        """Determine the primary document type."""
        extension = Path(filename).suffix.lower()
        
        for doc_type, extensions in self.supported_formats.items():
            if extension in extensions:
                return doc_type
                
        # Fallback to content type analysis
        if content_type:
            if content_type.startswith('image/'):
                return 'image'
            elif 'pdf' in content_type:
                return 'pdf'
            elif 'word' in content_type or 'document' in content_type:
                return 'word'
            elif 'csv' in content_type or 'text' in content_type:
                return 'csv'
                
        return 'unknown'
        
    async def _analyze_image_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze image document structure for OCR optimization."""
        try:
            # Load image
            image = cv2.imread(file_path)
            if image is None:
                return {"error": "Could not load image"}
                
            height, width = image.shape[:2]
            
            # Convert to grayscale for analysis
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Detect text regions using MSER
            mser = cv2.MSER_create()
            regions, _ = mser.detectRegions(gray)
            
            # Analyze image quality
            laplacian_var = cv2.Laplacian(gray, cv2.CV_64F).var()
            blur_score = "sharp" if laplacian_var > 100 else "moderate" if laplacian_var > 50 else "blurry"
            
            # Detect orientation
            contours, _ = cv2.findContours(gray, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # Calculate text density
            text_regions = len(regions)
            text_density = text_regions / (width * height / 10000)  # per 100x100 pixel block
            
            # Detect columns (simplified)
            columns = self._detect_columns(gray)
            
            # Language detection hints
            has_amharic_characteristics = self._detect_amharic_characteristics(gray)
            
            structure_info = {
                "width": width,
                "height": height,
                "aspect_ratio": round(width / height, 2),
                "text_regions_detected": text_regions,
                "text_density": round(text_density, 2),
                "image_quality": blur_score,
                "blur_score": round(laplacian_var, 2),
                "estimated_columns": columns,
                "has_amharic_features": has_amharic_characteristics,
                "color_mode": "color" if len(image.shape) == 3 else "grayscale",
                "recommended_dpi": self._recommend_dpi(width, height),
                "preprocessing_needed": blur_score == "blurry" or text_density < 0.5
            }
            
            return structure_info
            
        except Exception as e:
            logger.error(f"Image analysis failed: {e}")
            return {"error": str(e)}
            
    async def _analyze_pdf_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze PDF document structure."""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            
            structure_info = {
                "page_count": len(doc),
                "is_searchable": False,
                "has_images": False,
                "has_text": False,
                "pages_info": [],
                "text_extraction_possible": True,
                "estimated_complexity": "simple"
            }
            
            # Analyze first few pages in detail
            sample_pages = min(3, len(doc))
            
            for page_num in range(sample_pages):
                page = doc[page_num]
                
                # Check for text
                text = page.get_text()
                has_text = len(text.strip()) > 0
                
                # Check for images
                image_list = page.get_images()
                has_images = len(image_list) > 0
                
                # Get page dimensions
                rect = page.rect
                
                page_info = {
                    "page": page_num + 1,
                    "width": rect.width,
                    "height": rect.height,
                    "has_text": has_text,
                    "text_length": len(text),
                    "image_count": len(image_list),
                    "is_likely_scanned": has_images and not has_text
                }
                
                structure_info["pages_info"].append(page_info)
                
                # Update overall flags
                if has_text:
                    structure_info["has_text"] = True
                    structure_info["is_searchable"] = True
                if has_images:
                    structure_info["has_images"] = True
                    
            # Determine complexity
            scanned_pages = sum(1 for p in structure_info["pages_info"] if p["is_likely_scanned"])
            if scanned_pages > 0:
                structure_info["estimated_complexity"] = "complex" if scanned_pages > 1 else "moderate"
                
            doc.close()
            
            return structure_info
            
        except Exception as e:
            logger.error(f"PDF analysis failed: {e}")
            return {"error": str(e), "text_extraction_possible": False}
            
    async def _analyze_word_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze Word document structure."""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            
            # Count elements
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            # Analyze text content
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
                
            word_count = len(full_text.split())
            has_amharic = any('\u1200' <= char <= '\u137F' for char in full_text)
            
            # Check for images
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
                    
            structure_info = {
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "word_count": word_count,
                "character_count": len(full_text),
                "has_tables": table_count > 0,
                "has_images": image_count > 0,
                "image_count": image_count,
                "has_amharic_text": has_amharic,
                "estimated_complexity": "complex" if table_count > 5 or image_count > 3 else "simple",
                "text_extractable": True
            }
            
            return structure_info
            
        except Exception as e:
            logger.error(f"Word document analysis failed: {e}")
            return {"error": str(e), "text_extractable": False}
            
    async def _analyze_csv_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze CSV file structure."""
        try:
            import pandas as pd
            
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            df = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding, nrows=100)  # Sample first 100 rows
                    used_encoding = encoding
                    break
                except Exception:
                    continue
                    
            if df is None:
                return {"error": "Could not read CSV file"}
                
            # Analyze structure
            row_count = len(df)
            column_count = len(df.columns)
            
            # Detect data types
            column_info = {}
            for col in df.columns:
                col_data = df[col].dropna()
                if len(col_data) > 0:
                    # Check for Amharic text
                    has_amharic = any(
                        any('\u1200' <= char <= '\u137F' for char in str(val))
                        for val in col_data.head(10)
                    )
                    
                    column_info[col] = {
                        "dtype": str(df[col].dtype),
                        "null_count": df[col].isnull().sum(),
                        "has_amharic": has_amharic,
                        "sample_values": col_data.head(3).tolist()
                    }
                    
            structure_info = {
                "row_count": row_count,
                "column_count": column_count,
                "encoding": used_encoding,
                "column_info": column_info,
                "has_header": True,  # Assume first row is header
                "estimated_complexity": "complex" if column_count > 10 or row_count > 1000 else "simple",
                "processing_possible": True
            }
            
            return structure_info
            
        except Exception as e:
            logger.error(f"CSV analysis failed: {e}")
            return {"error": str(e), "processing_possible": False}
            
    def _detect_columns(self, gray_image: np.ndarray) -> int:
        """Detect number of text columns in an image."""
        try:
            # Simple column detection using vertical projection
            height, width = gray_image.shape
            vertical_projection = np.sum(gray_image < 128, axis=0)  # Count dark pixels
            
            # Smooth the projection
            from scipy import ndimage
            smoothed = ndimage.gaussian_filter1d(vertical_projection, sigma=width // 100)
            
            # Find valleys (potential column separators)
            threshold = np.mean(smoothed) * 0.3
            valleys = []
            
            for i in range(1, len(smoothed) - 1):
                if smoothed[i] < threshold and smoothed[i] < smoothed[i-1] and smoothed[i] < smoothed[i+1]:
                    valleys.append(i)
                    
            # Estimate columns (valleys + 1, but at least 1)
            return max(1, len(valleys) + 1) if valleys else 1
            
        except Exception:
            return 1  # Default to single column
            
    def _detect_amharic_characteristics(self, gray_image: np.ndarray) -> bool:
        """Detect visual characteristics that suggest Amharic text."""
        try:
            # Simple heuristic: Amharic text has specific character density patterns
            # This is a placeholder - real implementation would use more sophisticated methods
            
            # Calculate text region characteristics
            contours, _ = cv2.findContours(gray_image, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            if not contours:
                return False
                
            # Analyze character-like regions
            char_regions = 0
            for contour in contours:
                area = cv2.contourArea(contour)
                if 50 < area < 1000:  # Typical character size range
                    char_regions += 1
                    
            # Basic heuristic: if we have many small regions, might be Amharic
            return char_regions > 20
            
        except Exception:
            return False
            
    def _recommend_dpi(self, width: int, height: int) -> int:
        """Recommend optimal DPI for OCR based on image dimensions."""
        # Simple heuristic based on image size
        total_pixels = width * height
        
        if total_pixels > 8000000:  # Large image
            return 300
        elif total_pixels > 2000000:  # Medium image
            return 300
        else:  # Small image
            return 400  # Higher DPI for better quality
            
    def _generate_processing_recommendations(
        self,
        doc_type: str,
        structure_info: Dict[str, Any],
        basic_info: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate processing recommendations based on analysis."""
        
        recommendations = {
            "primary_processor": doc_type,
            "complexity": "simple",
            "estimated_time": 10,  # seconds
            "preprocessing_steps": [],
            "ocr_config": {},
            "quality_checks": [],
            "parallel_processing": False
        }
        
        if doc_type == "image":
            # Image-specific recommendations
            if structure_info.get("blur_score", 0) < 50:
                recommendations["preprocessing_steps"].append("image_enhancement")
                recommendations["complexity"] = "moderate"
                recommendations["estimated_time"] = 20
                
            if structure_info.get("estimated_columns", 1) > 1:
                recommendations["ocr_config"]["psm"] = 6  # Uniform block of text
                recommendations["complexity"] = "moderate"
                
            if structure_info.get("has_amharic_features"):
                recommendations["ocr_config"]["language"] = "amh"
                recommendations["quality_checks"].append("amharic_validation")
                
        elif doc_type == "pdf":
            page_count = structure_info.get("page_count", 1)
            
            if page_count > 10:
                recommendations["parallel_processing"] = True
                recommendations["complexity"] = "complex"
                recommendations["estimated_time"] = page_count * 3
                
            if not structure_info.get("is_searchable"):
                recommendations["preprocessing_steps"].append("pdf_to_image")
                recommendations["ocr_config"]["language"] = "amh+eng"
                
        elif doc_type == "word":
            if structure_info.get("has_tables"):
                recommendations["preprocessing_steps"].append("table_extraction")
                recommendations["complexity"] = "moderate"
                
        elif doc_type == "csv":
            if structure_info.get("column_count", 0) > 10:
                recommendations["complexity"] = "complex"
                recommendations["parallel_processing"] = True
                
        # Adjust time estimates based on file size
        size_mb = basic_info.get("file_size_mb", 0)
        if size_mb > 10:
            recommendations["estimated_time"] *= 2
            
        return recommendations
        
    def _calculate_confidence_scores(
        self,
        doc_type: str,
        structure_info: Dict[str, Any],
        basic_info: Dict[str, Any]
    ) -> Dict[str, float]:
        """Calculate confidence scores for various aspects."""
        
        confidence_scores = {
            "type_detection": 0.9,  # Default high confidence
            "structure_analysis": 0.8,
            "processing_success": 0.7,
            "quality_expectation": 0.8
        }
        
        # Adjust based on file characteristics
        if doc_type == "unknown":
            confidence_scores["type_detection"] = 0.3
            confidence_scores["processing_success"] = 0.4
            
        if "error" in structure_info:
            confidence_scores["structure_analysis"] = 0.2
            confidence_scores["processing_success"] = 0.3
            
        # Image-specific adjustments
        if doc_type == "image":
            blur_score = structure_info.get("blur_score", 100)
            if blur_score < 50:
                confidence_scores["quality_expectation"] = 0.4
            elif blur_score > 100:
                confidence_scores["quality_expectation"] = 0.95
                
        return confidence_scores
        
    def _detect_processing_issues(
        self,
        doc_type: str,
        structure_info: Dict[str, Any],
        basic_info: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Detect potential processing issues."""
        
        issues = []
        
        # File accessibility issues
        if not basic_info.get("is_readable"):
            issues.append({
                "type": "file_access",
                "severity": "critical",
                "description": "File is not readable or corrupted",
                "recommendation": "Re-upload the file"
            })
            
        # Size issues
        size_mb = basic_info.get("file_size_mb", 0)
        if size_mb > 100:
            issues.append({
                "type": "file_size",
                "severity": "warning",
                "description": f"Large file size ({size_mb}MB) may cause processing delays",
                "recommendation": "Consider splitting into smaller files"
            })
            
        # Format-specific issues
        if doc_type == "image":
            if structure_info.get("blur_score", 100) < 50:
                issues.append({
                    "type": "image_quality",
                    "severity": "high",
                    "description": "Image appears blurry, OCR accuracy may be reduced",
                    "recommendation": "Use higher quality scan if available"
                })
                
            if structure_info.get("text_density", 1) < 0.2:
                issues.append({
                    "type": "content_density",
                    "severity": "medium",
                    "description": "Low text density detected, document may be mostly graphics",
                    "recommendation": "Verify document contains extractable text"
                })
                
        elif doc_type == "pdf":
            if structure_info.get("page_count", 0) > 100:
                issues.append({
                    "type": "document_length",
                    "severity": "medium",
                    "description": "Very long document may require extended processing time",
                    "recommendation": "Consider processing in batches"
                })
                
        elif doc_type == "unknown":
            issues.append({
                "type": "format_detection",
                "severity": "high",
                "description": "Document format could not be reliably detected",
                "recommendation": "Verify file format and extension"
            })
            
        return issues


# Global document analyzer instance
_document_analyzer: Optional[DocumentAnalyzerAgent] = None


def get_document_analyzer(settings: Settings) -> DocumentAnalyzerAgent:
    """Get the global document analyzer agent instance."""
    global _document_analyzer
    if _document_analyzer is None:
        _document_analyzer = DocumentAnalyzerAgent(settings)
    return _document_analyzer